from flask import Flask, request, render_template, send_file
import os, pandas as pd, numpy as np
import plotly.express as px
import plotly.graph_objects as go
from modules.parser import load_deseq2
from modules.plots import volcano_plot, ma_plot, heatmap_top50
from modules.enrichment import run_enrichment
from modules.report import generate_report

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = "uploads"
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/analyze", methods=["POST"])
def analyze():
    if "file" not in request.files:
        return "No file uploaded", 400
    file = request.files["file"]
    if file.filename == "":
        return "No file selected", 400
    filepath = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
    file.save(filepath)
    df = load_deseq2(filepath)
    enrichment = run_enrichment(df)
    stats = {"total_genes": len(df), "significant": int(df["significant"].sum()),
        "upregulated": int((df["direction"]=="upregulated").sum()),
        "downregulated": int((df["direction"]=="downregulated").sum())}
    return render_template("results.html", stats=stats,
        volcano=volcano_plot(df), ma=ma_plot(df), heatmap=heatmap_top50(df),
        enrichment=enrichment, filepath=filepath)

@app.route("/download_report", methods=["POST"])
def download_report():
    filepath = request.form.get("filepath")
    df = load_deseq2(filepath)
    enrichment = run_enrichment(df)
    df2 = df.copy()
    df2["-log10padj"] = -np.log10(df2["padj"].clip(lower=1e-300))
    df2["log10baseMean"] = np.log10(df2["baseMean"].clip(lower=1e-10))
    color_map = {"upregulated":"red","downregulated":"blue","not significant":"grey"}
    v_fig = px.scatter(df2, x="log2FoldChange", y="-log10padj", color="direction",
        color_discrete_map=color_map, title="Volcano Plot")
    m_fig = px.scatter(df2, x="log10baseMean", y="log2FoldChange", color="direction",
        color_discrete_map=color_map, title="MA Plot")
    gene_col = "symbol" if "symbol" in df.columns else df.columns[0]
    top50 = pd.concat([df[df["direction"]=="upregulated"].nsmallest(25,"padj"),
        df[df["direction"]=="downregulated"].nsmallest(25,"padj")])
    h_fig = go.Figure(data=go.Heatmap(z=top50["log2FoldChange"].values.reshape(-1,1),
        y=top50[gene_col].values, x=["log2FC"], colorscale="RdBu_r", zmid=0))
    report_path = generate_report(df, enrichment, v_fig, m_fig, h_fig)
    return send_file(report_path, as_attachment=True, download_name="transcriptomic_report.pdf")


@app.route("/download_enrichment", methods=["POST"])
def download_enrichment():
    filepath = request.form.get("filepath")
    fmt = request.form.get("format")
    df = load_deseq2(filepath)
    enrichment = run_enrichment(df)
    clean = {k: v for k, v in enrichment.items() if k != "error"}
    
    if fmt == "csv":
        import io
        combined = []
        for key, table in clean.items():
            table = table.copy()
            table.insert(0, "section", key)
            combined.append(table)
        import pandas as pd
        out = pd.concat(combined)[["section","term","pvalue","adj_pvalue","combined_score"]]
        buf = io.StringIO()
        out.to_csv(buf, index=False)
        buf.seek(0)
        from flask import Response
        return Response(buf.getvalue(), mimetype="text/csv",
            headers={"Content-Disposition": "attachment;filename=enrichment_results.csv"})
    
    elif fmt == "docx":
        from docx import Document
        from docx.shared import Pt
        import io
        doc = Document()
        doc.add_heading("Enrichment Analysis Results", 0)
        for key, table in clean.items():
            doc.add_heading(key.replace("_", " ").title(), level=1)
            t = doc.add_table(rows=1, cols=4)
            t.style = "Table Grid"
            hdr = t.rows[0].cells
            for i, h in enumerate(["Term","P-value","Adj. P-value","Combined Score"]):
                hdr[i].text = h
            for _, row in table.head(15).iterrows():
                cells = t.add_row().cells
                cells[0].text = str(row["term"])
                cells[1].text = str(row["pvalue"])
                cells[2].text = str(row["adj_pvalue"])
                cells[3].text = str(row["combined_score"])
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf, as_attachment=True,
            download_name="enrichment_results.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


@app.route("/interpret", methods=["POST"])
def interpret():
    filepath = request.form.get("filepath")
    api_key = request.form.get("api_key") or os.environ.get("GROQ_API_KEY", "")
    context = request.form.get("context", "treatment vs control")
    df = load_deseq2(filepath)
    enrichment = run_enrichment(df)
    stats = {"total_genes": len(df), "significant": int(df["significant"].sum()),
        "upregulated": int((df["direction"]=="upregulated").sum()),
        "downregulated": int((df["direction"]=="downregulated").sum())}
    from modules.ai_interpret import get_biological_interpretation
    interpretation = get_biological_interpretation(stats, enrichment, api_key, context)
    return interpretation


@app.route("/download_section", methods=["POST"])
def download_section():
    filepath = request.form.get("filepath")
    section = request.form.get("section")
    fmt = request.form.get("format")
    df = load_deseq2(filepath)
    enrichment = run_enrichment(df)
    
    if section not in enrichment or section == "error":
        return "Section not found", 404
    
    table = enrichment[section]
    
    if fmt == "csv":
        import io
        buf = io.StringIO()
        table[["term","pvalue","adj_pvalue","combined_score"]].to_csv(buf, index=False)
        buf.seek(0)
        from flask import Response
        return Response(buf.getvalue(), mimetype="text/csv",
            headers={"Content-Disposition": f"attachment;filename={section}.csv"})
    
    elif fmt == "docx":
        from docx import Document
        import io
        doc = Document()
        doc.add_heading(section.replace("_", " ").title(), 0)
        t = doc.add_table(rows=1, cols=4)
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        for i, h in enumerate(["Term","P-value","Adj. P-value","Combined Score"]):
            hdr[i].text = h
        for _, row in table.head(15).iterrows():
            cells = t.add_row().cells
            cells[0].text = str(row["term"])
            cells[1].text = str(row["pvalue"])
            cells[2].text = str(row["adj_pvalue"])
            cells[3].text = str(row["combined_score"])
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf, as_attachment=True,
            download_name=f"{section}.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=7860)
